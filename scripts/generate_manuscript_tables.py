"""Generate manuscript-ready RVPI tables, captions, CSVs, Markdown and DOCX.

All language treats RVPI as RTLS value potential / potential-value
prioritisation. No table estimates causal RTLS effects, savings, or ROI.
"""
from __future__ import annotations

import json
import math
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
FINAL = ROOT / "data" / "final"
TABLES = ROOT / "outputs" / "tables"
DOCX_PATH = TABLES / "manuscript_tables_and_figure_captions.docx"
MD_PATH = TABLES / "manuscript_tables_and_figure_captions.md"

BLUE = "1F4E78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F7F9FB"
INK = RGBColor(31, 31, 31)


def fnum(value: object, decimals: int = 2) -> str:
    if pd.isna(value):
        return "-"
    number = float(value)
    # Suppress the visually distracting "-0.00" produced by floating-point
    # noise after within-sample standardisation.
    if abs(number) < 0.5 * (10 ** -decimals):
        number = 0.0
    return f"{number:,.{decimals}f}"


def dataframe_markdown(frame: pd.DataFrame) -> str:
    clean = frame.astype(object).where(pd.notna(frame), "")
    header = "| " + " | ".join(map(str, clean.columns)) + " |"
    divider = "|" + "|".join(["---"] * len(clean.columns)) + "|"
    body = ["| " + " | ".join(str(v).replace("|", "\\|") for v in row) + " |"
            for row in clean.itertuples(index=False, name=None)]
    return "\n".join([header, divider, *body])


def build_tables() -> dict[str, pd.DataFrame]:
    primary = pd.read_csv(FINAL / "rvpi_primary_acute_trusts.csv")
    full = pd.read_csv(FINAL / "rvpi_full_provider_sensitivity.csv")
    flow = pd.read_csv(TABLES / "primary_sample_flow.csv")
    comparison = pd.read_csv(TABLES / "sensitivity_comparison.csv")

    table1 = pd.DataFrame([
        ["Bed occupancy pressure", "NHS England KH03 overnight beds, Q4 2025-26",
         "General & Acute occupied beds / General & Acute available beds", "Higher = greater occupancy pressure"],
        ["A&E delay burden", "NHS England monthly A&E, March 2026",
         "1 - attendances within 4 hours / total attendances", "Higher = greater four-hour delay burden"],
        ["Emergency admissions intensity", "NHS England monthly A&E, March 2026; KH03 Q4 2025-26",
         "Emergency admissions / General & Acute available beds", "Higher = greater emergency throughput per bed"],
        ["Waiting-list burden", "NHS England revised provider RTT, March 2026",
         "Incomplete pathways over 18 weeks / total incomplete pathways", "Higher = greater long-wait burden"],
        ["Estates burden", "NHS England ERIC site data, 2024-25",
         "Total cost to eradicate backlog / occupied floor area (GBP/m2)", "Higher = greater estates backlog burden"],
        ["Resource/capacity strain", "NHS England monthly A&E, March 2026; KH03 Q4 2025-26",
         "A&E attendances / General & Acute available beds", "Higher = greater activity relative to capacity"],
        ["RVPI", "All six component sources",
         "Equal-weight mean of available component z-scores; minimum 4 of 6", "Higher = greater RTLS value potential"],
    ], columns=["Indicator", "Data source and period", "Construction", "Direction"])

    table2 = flow[["step", "criterion", "sequential_excluded", "remaining"]].copy()
    table2.columns = ["Stage", "Criterion", "Excluded at stage", "Remaining"]
    labels = {
        "full_provider_spine": "Full provider spine",
        "exclude_non_general_acute_type": "Exclude non-general-acute type",
        "exclude_missing_general_acute_beds": "Exclude missing General & Acute beds",
        "exclude_zero_or_negative_general_acute_beds": "Exclude non-positive General & Acute beds",
        "exclude_missing_ae_activity": "Exclude missing A&E activity",
        "exclude_zero_or_negative_ae_activity": "Exclude non-positive A&E activity",
        "exclude_fewer_than_four_components": "Exclude <4 available components",
        "primary_acute_trust_sample": "Primary acute-trust sample",
    }
    table2["Stage"] = table2.Stage.map(labels).fillna(table2.Stage)

    specs = [
        ("Bed occupancy pressure", "bed_occupancy_pressure", "%", 100),
        ("A&E delay burden", "ae_delay_burden", "%", 100),
        ("Emergency admissions intensity", "emergency_admissions_intensity", "admissions per bed", 1),
        ("Waiting-list burden", "waiting_list_burden", "%", 100),
        ("Estates burden", "estates_burden", "GBP per m2", 1),
        ("Resource/capacity strain", "resource_capacity_strain", "attendances per bed", 1),
        ("RVPI", "rvpi_z", "mean z-score", 1),
    ]
    rows = []
    for label, column, unit, scale in specs:
        series = pd.to_numeric(primary[column], errors="coerce").dropna() * scale
        rows.append([label, unit, len(series), series.mean(), series.std(ddof=1), series.median(),
                     series.quantile(.25), series.quantile(.75), series.min(), series.max()])
    table3 = pd.DataFrame(rows, columns=["Measure", "Unit", "N", "Mean", "SD", "Median", "P25", "P75", "Min", "Max"])
    for col in ["Mean", "SD", "Median", "P25", "P75", "Min", "Max"]:
        table3[col] = table3[col].map(lambda x: fnum(x, 2))

    ranked = primary.sort_values("rvpi_rank").copy()
    top_decile_n = math.ceil(len(primary) * .10)
    ranking_cols = ["rvpi_rank", "trust_code", "trust_name", "provider_type_source",
                    "rvpi_z", "rvpi_percentile", "cluster", "components_available"]
    table4 = ranked.head(top_decile_n)[ranking_cols].copy()
    s1 = ranked.head(25)[ranking_cols].copy()
    ranking_names = ["Rank", "ODS code", "Trust", "ERIC trust type", "RVPI", "Percentile", "Cluster", "Components"]
    table4.columns = ranking_names
    s1.columns = ranking_names
    for frame in (table4, s1):
        frame["Rank"] = frame["Rank"].astype(int)
        frame["RVPI"] = frame["RVPI"].map(lambda x: fnum(x, 3))
        frame["Percentile"] = frame["Percentile"].map(lambda x: fnum(x, 1))
        frame["Cluster"] = frame["Cluster"].astype("Int64")
        frame["Components"] = frame["Components"].astype(int)

    s2 = comparison[["trust_code", "trust_name", "primary_rvpi_rank", "full_rvpi_rank",
                     "rank_change_primary_minus_full", "primary_rvpi_percentile",
                     "full_rvpi_percentile", "primary_cluster", "full_cluster"]].copy()
    s2.columns = ["ODS code", "Trust", "Primary rank", "Full-provider rank", "Rank difference",
                  "Primary percentile", "Full-provider percentile", "Primary cluster", "Full-provider cluster"]
    for col in ["Primary rank", "Full-provider rank", "Rank difference", "Primary cluster", "Full-provider cluster"]:
        s2[col] = s2[col].astype("Int64")
    for col in ["Primary percentile", "Full-provider percentile"]:
        s2[col] = s2[col].map(lambda x: fnum(x, 1))
    s2 = s2.sort_values("Primary rank")

    return {"table1": table1, "table2": table2, "table3": table3, "table4": table4, "s1": s1, "s2": s2}


def write_csvs(tables: dict[str, pd.DataFrame]) -> None:
    names = {
        "table1": "manuscript_table_1_data_sources_indicators.csv",
        "table2": "manuscript_table_2_primary_sample_flow.csv",
        "table3": "manuscript_table_3_descriptive_statistics.csv",
        "table4": "manuscript_table_4_top_decile_rvpi.csv",
        "s1": "supplementary_table_s1_top25_rankings.csv",
        "s2": "supplementary_table_s2_sensitivity_comparison.csv",
    }
    TABLES.mkdir(parents=True, exist_ok=True)
    for key, filename in names.items():
        tables[key].to_csv(TABLES / filename, index=False)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge, val, size, color in (
        ("top", "single", "8", BLUE), ("bottom", "single", "8", BLUE),
        ("insideH", "single", "3", "B7C5D3"),
        ("left", "nil", "0", "FFFFFF"), ("right", "nil", "0", "FFFFFF"),
        ("insideV", "nil", "0", "FFFFFF"),
    ):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), val)
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)
        borders.append(tag)


def set_fixed_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_table(doc: Document, frame: pd.DataFrame, widths: list[float], font_size=8.5,
              left_columns: set[int] | None = None) -> None:
    left_columns = left_columns or {0}
    table = doc.add_table(rows=1, cols=len(frame.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_fixed_layout(table)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_header(header)
    for j, label in enumerate(frame.columns):
        cell = header.cells[j]
        cell.width = Inches(widths[j])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j in left_columns else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(label))
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(BLUE)
    for row_index, values in enumerate(frame.itertuples(index=False, name=None)):
        row = table.add_row()
        prevent_row_split(row)
        if row_index % 2:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_FILL)
        for j, value in enumerate(values):
            cell = row.cells[j]
            cell.width = Inches(widths[j])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j in left_columns else WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run("-" if pd.isna(value) else str(value))
            run.font.size = Pt(font_size)
            run.font.color.rgb = INK


def add_caption(doc: Document, label: str, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"{label}. {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(BLUE)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    lead = p.add_run("Note. ")
    lead.bold = True
    lead.font.size = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(8)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar"); fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar"); fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Inches(11), Inches(8.5)
    section.top_margin = Inches(.55); section.bottom_margin = Inches(.55)
    section.left_margin = Inches(.55); section.right_margin = Inches(.55)
    section.header_distance = Inches(.3); section.footer_distance = Inches(.3)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(10); normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.15
    for name, size, color in (("Title", 18, BLUE), ("Heading 1", 14, BLUE), ("Heading 2", 12, BLUE)):
        style = styles[name]
        style.font.name = "Calibri"; style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color)
    footer = section.footer
    p = footer.paragraphs[0]
    p.add_run("RVPI manuscript tables  |  ").font.size = Pt(8)
    add_page_number(p)
    return doc


def write_docx(tables: dict[str, pd.DataFrame], captions: dict[str, str]) -> None:
    doc = configure_document()
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Publication-ready RVPI tables and figure captions")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Primary acute-trust analysis and full-provider sensitivity analysis").italic = True
    caveat = doc.add_paragraph()
    caveat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caveat.add_run("RVPI measures RTLS value potential for potential-value prioritisation; it is not a causal RTLS effect or ROI estimate.").bold = True

    blocks = [
        ("Table 1", "Data sources and RVPI indicators", "table1", [1.45, 2.35, 3.75, 2.3], 8.2, {0, 1, 2, 3},
         "KH03, A&E and revised provider RTT are aligned to 2025-26; ERIC is 2024-25. ODS organisation codes are used for linkage without fuzzy matching."),
        ("Table 2", "Primary acute-trust sample flow", "table2", [2.25, 4.9, 1.25, 1.25], 8.5, {0, 1},
         "Exclusions are sequential. Across the full spine, overlapping flags identified 35 providers with zero General & Acute beds and 9 with zero A&E activity."),
        ("Table 3", "Descriptive statistics for RVPI components and RVPI in the primary acute-trust sample", "table3",
         [1.8, 1.15, .45, .75, .75, .75, .75, .75, .75, .75], 8.0, {0, 1},
         "N=114 unless missing. Percentage indicators are shown in percentage points. RVPI is the equal-weight mean of available component z-scores within the primary sample."),
        ("Table 4", "Top-decile acute trusts by RTLS value potential", "table4",
         [.45, .65, 3.15, 1.45, .6, .75, .55, .65], 7.8, {1, 2, 3},
         "The empirical top decile is defined as the first ceiling(0.10 x 114)=12 ranked trusts. A high RVPI denotes greater RTLS value potential, not realised benefit, causal impact, or ROI."),
        ("Supplementary Table S1", "Top 25 primary acute-trust RVPI rankings", "s1",
         [.45, .65, 3.15, 1.45, .6, .75, .55, .65], 7.6, {1, 2, 3},
         "Scores, percentiles and clusters are standardised within the primary acute-trust sample. Rankings support potential-value prioritisation only."),
        ("Supplementary Table S2", "Sensitivity comparison of primary and full-provider RVPI rankings", "s2",
         [.6, 3.1, .65, .8, .75, .8, .85, .75, .85], 7.2, {0, 1},
         "Rank difference = primary rank minus full-provider rank; negative values indicate a higher position in the primary analysis. Spearman rho=0.879; median absolute rank change=7.5. Cluster solutions were not stable (primary k=5, full-provider k=2, adjusted Rand index=0.000)."),
    ]
    for label, title_text, key, widths, size, left, note in blocks:
        doc.add_page_break()
        if key != "s2":
            add_caption(doc, label, title_text)
            add_table(doc, tables[key], widths, font_size=size, left_columns=left)
            add_note(doc, note)
            continue

        # Word occasionally omits a repeated header on the final continuation
        # page of a very long table. Explicit pagination guarantees that every
        # manuscript page carries both the S2 label and column headings.
        # Twenty-seven rows leave reliable room for the continuation heading;
        # the short final page also accommodates the sensitivity-analysis note.
        page_size = 27
        pages = [tables[key].iloc[start:start + page_size]
                 for start in range(0, len(tables[key]), page_size)]
        for page_index, page in enumerate(pages):
            if page_index:
                doc.add_page_break()
            continued = " (continued)" if page_index else ""
            add_caption(doc, f"{label}{continued}", title_text)
            add_table(doc, page, widths, font_size=size, left_columns=left)
            if page_index == len(pages) - 1:
                add_note(doc, note)

    doc.add_page_break()
    add_caption(doc, "Figure captions", "Manuscript-ready text")
    for key in ("distribution", "clusters"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.add_run(captions[key]).font.size = Pt(10)
    doc.save(DOCX_PATH)


def write_markdown(tables: dict[str, pd.DataFrame], captions: dict[str, str]) -> None:
    titles = {
        "table1": "Table 1. Data sources and RVPI indicators",
        "table2": "Table 2. Primary acute-trust sample flow",
        "table3": "Table 3. Descriptive statistics for the six RVPI components and RVPI",
        "table4": "Table 4. Top-decile acute trusts by RTLS value potential",
        "s1": "Supplementary Table S1. Top 25 primary acute-trust RVPI rankings",
        "s2": "Supplementary Table S2. Sensitivity comparison between primary and full-provider RVPI",
    }
    parts = ["# Manuscript tables and figure captions", "",
             "RVPI supports RTLS value potential and potential-value prioritisation; it is not a causal RTLS effect or ROI estimate."]
    for key in ("table1", "table2", "table3", "table4", "s1", "s2"):
        parts.extend(["", f"## {titles[key]}", "", dataframe_markdown(tables[key])])
    parts.extend(["", "## Figure captions", "", captions["distribution"], "", captions["clusters"]])
    MD_PATH.write_text("\n".join(parts) + "\n", encoding="utf-8")


def main() -> None:
    tables = build_tables()
    captions = {
        "distribution": (
            "Figure 1. Distribution of the Real-Time Location System Value Potential Index (RVPI) in the "
            "primary acute-trust sample (n=114). RVPI is the equal-weight mean of up to six standardised "
            "operational-pressure components, with at least four required; the dashed vertical line marks the "
            "sample mean of zero. Higher values indicate greater RTLS value potential for potential-value "
            "prioritisation and do not represent causal RTLS effects, realised savings, or return on investment."
        ),
        "clusters": (
            "Figure 2. Exploratory K-means clusters in the primary acute-trust sample. Points show the 112 trusts "
            "with complete data for all six RVPI components; the x-axis is General & Acute bed occupancy, the "
            "y-axis is A&E four-hour delay burden, point size reflects the primary-sample RVPI percentile, and "
            "colour denotes the five-cluster solution selected by the silhouette criterion. Clusters describe "
            "multivariate operational profiles for RTLS value potential and are not treatment-effect groups, "
            "causal estimates, or ROI categories."
        ),
    }
    write_csvs(tables)
    (TABLES / "figure_captions.txt").write_text(captions["distribution"] + "\n\n" + captions["clusters"] + "\n", encoding="utf-8")
    write_markdown(tables, captions)
    write_docx(tables, captions)
    print(f"Generated six manuscript tables and captions; DOCX: {DOCX_PATH}")


if __name__ == "__main__":
    main()
